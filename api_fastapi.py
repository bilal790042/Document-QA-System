"""
FastAPI Backend with File Upload Support

This API supports:
- Asking questions (existing functionality)
- Uploading documents (NEW: txt, docx, pdf)

Installation:
    pip install fastapi uvicorn python-multipart python-docx PyPDF2

Run:
    uvicorn api_with_upload:app --reload
"""

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Dict
import os
import tempfile
from pathlib import Path

# Import your existing services
from services.qa_service import QAService

# For file processing
try:
    from PyPDF2 import PdfReader
    from docx import Document
except ImportError:
    print("⚠️  Warning: Install PyPDF2 and python-docx for full file support")
    print("   pip install PyPDF2 python-docx")

# Initialize FastAPI app
app = FastAPI(title="QA Service API with File Upload")

# Enable CORS for React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize QA service once (singleton)
try:
    qa_service = QAService()
    print("✅ QA Service initialized successfully")
except Exception as e:
    print(f"❌ Error initializing QA Service: {e}")
    qa_service = None


# Request/Response models
class QuestionRequest(BaseModel):
    question: str


class Source(BaseModel):
    source: str
    content: str


class AnswerResponse(BaseModel):
    answer: str
    sources: List[Source]
    question: str


# File processing functions
def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try different encoding if utf-8 fails
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")


def process_uploaded_file(file_path: str, filename: str) -> str:
    """Process uploaded file based on extension"""
    ext = Path(filename).suffix.lower()
    
    if ext == '.txt':
        return extract_text_from_txt(file_path)
    elif ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# API Endpoints
@app.get("/")
async def root():
    """Root endpoint - shows API is running"""
    return {
        "message": "QA Service API with File Upload",
        "endpoints": {
            "health": "/api/health/",
            "ask": "/api/ask/ (POST)",
            "upload": "/api/upload/ (POST)"
        }
    }


@app.get("/api/health/")
async def health_check():
    """Health check endpoint"""
    if qa_service is None:
        raise HTTPException(status_code=503, detail="QA Service not initialized")
    return {"status": "ok", "service": "ready"}


@app.post("/api/ask/", response_model=AnswerResponse)
async def ask_question(request: QuestionRequest):
    """
    Ask a question and get an AI-generated answer with sources
    
    Args:
        request: QuestionRequest with 'question' field
        
    Returns:
        AnswerResponse with answer, sources, and original question
    """
    if qa_service is None:
        raise HTTPException(
            status_code=503, 
            detail="QA Service is not initialized. Check your configuration."
        )
    
    if not request.question or not request.question.strip():
        raise HTTPException(
            status_code=400, 
            detail="Question cannot be empty"
        )
    
    try:
        # Get answer from QA service
        answer, docs = qa_service.ask(request.question)
        
        # Format sources
        sources = []
        for doc in docs:
            source_info = Source(
                source=doc.metadata.get("source", "unknown"),
                content=doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content
            )
            sources.append(source_info)
        
        return AnswerResponse(
            answer=answer,
            sources=sources,
            question=request.question
        )
    
    except Exception as e:
        print(f"Error processing question: {e}")
        raise HTTPException(
            status_code=500, 
            detail=f"Error processing question: {str(e)}"
        )


@app.post("/api/upload/")
async def upload_files(files: List[UploadFile] = File(...)):
    """
    Upload one or more documents (txt, docx, pdf)
    Files will be processed and added to the vector store
    
    Args:
        files: List of uploaded files
        
    Returns:
        Success message with processed file names
    """
    if qa_service is None:
        raise HTTPException(
            status_code=503,
            detail="QA Service is not initialized"
        )
    
    if not files:
        raise HTTPException(
            status_code=400,
            detail="No files provided"
        )
    
    processed_files = []
    errors = []
    
    # Create temp directory for uploaded files
    temp_dir = tempfile.mkdtemp()
    
    try:
        for file in files:
            # Validate file type
            ext = Path(file.filename).suffix.lower()
            if ext not in ['.txt', '.pdf', '.docx', '.doc']:
                errors.append(f"{file.filename}: Unsupported file type")
                continue
            
            # Save uploaded file temporarily
            temp_file_path = os.path.join(temp_dir, file.filename)
            
            try:
                # Save file
                with open(temp_file_path, 'wb') as f:
                    content = await file.read()
                    f.write(content)
                
                # Extract text from file
                text_content = process_uploaded_file(temp_file_path, file.filename)
                
                if not text_content.strip():
                    errors.append(f"{file.filename}: No text content found")
                    continue
                
                # Add to vector store (you'll need to implement this in your QAService)
                # This assumes your QAService has an add_document method
                # Modify this based on your actual implementation
                if hasattr(qa_service, 'add_document'):
                    qa_service.add_document(text_content, {"source": file.filename})
                elif hasattr(qa_service, 'vectorstore'):
                    # Alternative: directly add to vectorstore
                    from langchain.schema import Document
                    doc = Document(page_content=text_content, metadata={"source": file.filename})
                    qa_service.vectorstore.add_documents([doc])
                else:
                    raise Exception("QA Service doesn't support adding documents")
                
                processed_files.append(file.filename)
                
            except Exception as e:
                errors.append(f"{file.filename}: {str(e)}")
                continue
            finally:
                # Clean up temp file
                if os.path.exists(temp_file_path):
                    os.remove(temp_file_path)
        
        # Clean up temp directory
        os.rmdir(temp_dir)
        
        if not processed_files:
            raise HTTPException(
                status_code=400,
                detail=f"No files were processed successfully. Errors: {'; '.join(errors)}"
            )
        
        response = {
            "message": "Files processed successfully",
            "files_processed": len(processed_files),
            "filenames": processed_files
        }
        
        if errors:
            response["warnings"] = errors
        
        return response
    
    except Exception as e:
        # Clean up on error
        if os.path.exists(temp_dir):
            import shutil
            shutil.rmtree(temp_dir)
        
        raise HTTPException(
            status_code=500,
            detail=f"Error processing files: {str(e)}"
        )


@app.get("/api/stats/")
async def get_stats():
    """Get basic statistics about the service"""
    return {
        "status": "operational",
        "version": "2.0.0",
        "features": ["qa", "file_upload"],
        "supported_formats": ["txt", "pdf", "docx"],
        "qa_service_ready": qa_service is not None
    }


if __name__ == "__main__":
    import uvicorn
    print("🚀 Starting FastAPI server with file upload support...")
    print("📝 API docs available at: http://localhost:8000/docs")
    print("📁 Upload endpoint: http://localhost:8000/api/upload/")
    print("💡 Frontend should connect to: http://localhost:8000")
    uvicorn.run(app, host="0.0.0.0", port=8000)